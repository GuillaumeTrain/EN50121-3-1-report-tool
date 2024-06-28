import sys

import pyqtgraph
from docx import Document
from enum import Enum
from functools import reduce
from pprint import pprint
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import pandas as pd
from PySide6.QtCore import QAbstractTableModel, Qt, QModelIndex, QAbstractItemModel
from PySide6.QtGui import QColor, QIntValidator, QStandardItemModel, QStandardItem, QBrush
from PySide6.QtWidgets import QWidget, QLineEdit, QLabel, QPushButton, QHBoxLayout, QVBoxLayout, QTableView, \
    QColorDialog, QTreeView, QSpinBox, QComboBox, QMessageBox, QListView, QFileSystemModel, QAbstractItemView
from numpy import ndarray
from docx.shared import RGBColor, Cm
import filebrowzer
from csvreader import FileReader

if 'numpy' not in sys.modules:
    import numpy as np
else:
    np = sys.modules['numpy']


from pyqtgraph.exporters import ImageExporter
import filebrowzer


class AmbientProcessor:
    def __init__(self, qamb_list: [filebrowzer.QListPlotItems], root: QWidget):
        self.root = root
        self.ambiant_filepath_list = qamb_list.get_filepath_list()
        # ambiant_filepaths, color = self.ambiant_list.get_item_color(0)
        self.frequency_axis: ndarray = None
        self.ambient_df_list: list[Meas] = []
        self.qlist = qamb_list
        try:
            self.color: QColor = qamb_list.get_item_color(0)
        except Exception:
            print("aucun ambient n'a été donné en donnée d'entrée ")
            self.color: QColor = QColor(92, 85, 233)
        self.name = "Max Ambient"
        self.merged_ambient: Meas = None
        self.get_data_from_csv()


    def get_merged_ambient_peaks(self, threshold=6):
        df = self.merged_ambient.data
        # print(df.values)
        # Trouver les pics
        threshold = threshold
        # df = self.merged_ambient.data
        all_frequencies = df['Frequency'].values
        merged_ambient = df['Field'].values
        self.peak_lower_base_index = 0
        self.peak_index = 0
        self.peak_upper_base_index = 0
        self.fisrtpeak_occured = False
        self.ambient_field_peaks = []

        for i in range(len(all_frequencies)):
            print(f"--------------------------------------\n")
            print(f"i = {i}")
            print(f"lowindex :{self.peak_lower_base_index} lowfreq = {all_frequencies[self.peak_lower_base_index]} ")
            print(f" pkindex : {self.peak_index} pkfreq= {all_frequencies[self.peak_index]}")
            print(f"lowindex :{self.peak_upper_base_index} lowfreq = {all_frequencies[self.peak_upper_base_index]} ")
            print(f"deltalow :{merged_ambient[self.peak_index] - merged_ambient[self.peak_lower_base_index]}")
            print(f"deltaup :{merged_ambient[self.peak_index] - merged_ambient[self.peak_upper_base_index]}")
            if i == 0:
                self.peak_lower_base_index = 0
                self.peak_index = 0
                self.peak_upper_base_index = 0
            elif merged_ambient[i] > merged_ambient[i - 1]:
                self.peak_upper_base_index = i
                if merged_ambient[i] > merged_ambient[self.peak_index]:
                    self.peak_index = i

                while ((merged_ambient[self.peak_index] - merged_ambient[
                    self.peak_lower_base_index + 1]) > threshold) and (
                        (self.peak_lower_base_index + 1) < self.peak_index):
                    self.peak_lower_base_index = self.peak_lower_base_index + 1

            elif merged_ambient[i] <= merged_ambient[i - 1]:
                self.peak_upper_base_index = i

                lessthan_lower = merged_ambient[i] < merged_ambient[self.peak_lower_base_index]
                upper_six_dB_frompeak = merged_ambient[self.peak_index] - merged_ambient[
                    self.peak_upper_base_index] > threshold
                lower_six_dB_frompeak = merged_ambient[self.peak_index] - merged_ambient[
                    self.peak_lower_base_index] > threshold

                if (lessthan_lower and lower_six_dB_frompeak) or (upper_six_dB_frompeak and not self.fisrtpeak_occured):

                    peak_freq = all_frequencies[self.peak_index]
                    peak_lvl = merged_ambient[self.peak_index]
                    peak_lower_base_freq = all_frequencies[self.peak_lower_base_index]
                    peak_lower_base_lvl = merged_ambient[self.peak_lower_base_index]
                    peak_upper_base_freq = all_frequencies[self.peak_upper_base_index]
                    peak_upper_base_lvl = merged_ambient[self.peak_upper_base_index]
                    self.ambient_field_peaks.append({
                        'Peak Index': self.peak_index,
                        'Peak Frequency': peak_freq,
                        'Peak Level': peak_lvl,
                        'Lower Base Index': self.peak_lower_base_index,
                        'Lower Base Frequency': peak_lower_base_freq,
                        'Lower Base Level': peak_lower_base_lvl,
                        'Upper Base Index': self.peak_upper_base_index,
                        'Upper Base Frequency': peak_upper_base_freq,
                        'Upper Base Level': peak_upper_base_lvl
                    })
                    self.peak_lower_base_index = i
                    self.peak_index = i
                    self.peak_upper_base_index = i
                    self.fisrtpeak_occured = True
                elif merged_ambient[i] < merged_ambient[self.peak_lower_base_index] and self.fisrtpeak_occured:
                    self.peak_lower_base_index = i
                    self.peak_index = i

        peaks_df = pd.DataFrame(self.ambient_field_peaks)
        print(peaks_df)

        targetsink: AmbientPostProcessor = None
        targetsink = self.root.ambient_postprocessor

        targetsink.set_peaks_df(peaks_df)
        return pd.DataFrame(peaks_df)

    def get_data_from_csv(self):
        # Créer un DataFrame pour chaque fichier, avec les colonnes 'Frequency', 'Field'
        item: filebrowzer.QMeasItem = None  # juste pour préconditionner le type et avec acces aux méthodes et attribut de la classe dans le code
        for item in self.qlist.items:
            f = FileReader()
            f.filepath = item.get_filepath()
            trace_nbr = int(item.trace_nbr.itemText(item.trace_nbr.currentIndex()))
            datas = f.get_arrayfromcsv(trace_nbr)
            meas = Meas(pd.DataFrame(datas, columns=['Frequency', 'Field']), self.color, item.get_filename(),
                        item.filepath)
            print(meas)
            self.ambient_df_list.append(meas)

    def set_frequency_axis(self, frequency_axis: ndarray = None):
        self.frequency_axis = frequency_axis
        if frequency_axis is None:
            self.frequency_axis = reduce(np.union1d, (meas.data['Frequency'] for meas in self.ambient_df_list))

    def interp_data(self):
        if self.ambient_df_list == []:
            print("faking ambient")
            fakeamb = Meas(pd.DataFrame([[self.frequency_axis.min(), -100], [self.frequency_axis.max(), -100]],
                                        columns=['Frequency', 'Field']), self.color, "defaultamb", "")
            self.ambient_df_list: [Meas] = [fakeamb]
        # Faire une interpolation sur une grille de fréquences commune, même si toutes les meas ont le même ensemble de fréquences
        for i, meas in enumerate(self.ambient_df_list):
            # Créer un nouveau DataFrame avec les valeurs interpolées
            interpolated_values = np.interp(self.frequency_axis, meas.data['Frequency'], meas.data['Field'])
            # Remplacer les valeurs en dehors de la plage d'interpolation par -100 dBµV/m
            interpolated_values = np.where(
                (self.frequency_axis < meas.data['Frequency'].min()) | (
                        self.frequency_axis > meas.data['Frequency'].max()), -100,
                interpolated_values)
            # Créer une nouvelle instance de Meas
            self.ambient_df_list[i] = Meas(pd.DataFrame({
                'Frequency': self.frequency_axis,
                'Field': interpolated_values,
            }), meas.color, meas.name, meas.filepath)

    def get_merged_ambient(self):
        return self.merged_ambient

    def merge_ambients(self):
        concatenated_data = pd.concat([amb.data for amb in self.ambient_df_list])
        # Group by 'Frequency' and keep the max 'Field' value
        grouped_data = concatenated_data.groupby('Frequency')['Field'].max().reset_index()
        # Create a new Meas object with the merged data
        self.merged_ambient = Meas(grouped_data, self.color, self.name, "")


class Meas:
    def __init__(self, data, color: QColor, name, filepath):
        self.data = pd.DataFrame(data, columns=['Frequency', 'Field'])
        self.color: QColor = color
        self.name: str = name
        self.filepath: str = filepath


class MeasProcessor:
    def __init__(self, qmeas_list: [filebrowzer.QListPlotItems]):
        self.frequency_axis: ndarray = None
        self.meas_list: list[Meas] = []
        self.qlist: filebrowzer.QListPlotItems = qmeas_list
        self.get_data_from_csv()
        # self.process_meas_list()

    def get_data_from_csv(self):
        # Créer un DataFrame pour chaque fichier, avec les colonnes 'Frequency', 'Field'
        item: filebrowzer.QMeasItem = None  # juste pour préconditionner le type et avec acces aux méthodes et attribut de la classe dans le code
        for item in self.qlist.items:
            f = FileReader()
            f.filepath = item.get_filepath()
            trace_nbr = int(item.trace_nbr.itemText(item.trace_nbr.currentIndex()))
            datas = f.get_arrayfromcsv(trace_nbr)
            meas = Meas(pd.DataFrame(datas, columns=['Frequency', 'Field']), item.get_color(), item.get_filename(),
                        item.filepath)
            self.meas_list.append(meas)

    def set_frequency_axis(self, frequency_axis: ndarray = None):
        self.frequency_axis = frequency_axis
        if frequency_axis is None:
            self.frequency_axis = reduce(np.union1d, (meas.data['Frequency'] for meas in self.meas_list))

    def interp_data(self):
        # Faire une interpolation sur une grille de fréquences commune, même si toutes les meas ont le même ensemble de fréquences
        for i, meas in enumerate(self.meas_list):
            # Créer un nouveau DataFrame avec les valeurs interpolées
            interpolated_values = np.interp(self.frequency_axis, meas.data['Frequency'], meas.data['Field'])
            # Remplacer les valeurs en dehors de la plage d'interpolation par -100 dBµV/m
            interpolated_values = np.where(
                (self.frequency_axis < meas.data['Frequency'].min()) | (
                        self.frequency_axis > meas.data['Frequency'].max()), -100,
                interpolated_values)
            # Créer une nouvelle instance de Meas
            self.meas_list[i] = Meas(pd.DataFrame({
                'Frequency': self.frequency_axis,
                'Field': interpolated_values,
            }), meas.color, meas.name, meas.filepath)

    def get_meas_list(self):
        return self.meas_list


class LimitProcessor:
    def __init__(self, qmeas_list: [filebrowzer.QListPlotItems]):
        self.frequency_axis: ndarray = None
        self.lim_list: list[Meas] = []
        self.qlist: filebrowzer.QListPlotItems = qmeas_list
        self.get_data_from_csv()
        # self.process_meas_list()

    def set_frequency_axis(self, frequency_axis: ndarray = None):
        self.frequency_axis = frequency_axis
        if frequency_axis is None:
            self.frequency_axis = reduce(np.union1d, (lim.data['Frequency'] for lim in self.lim_list))

    def get_data_from_csv(self):
        # Créer un DataFrame pour chaque fichier, avec les colonnes 'Frequency', 'Field'
        item: filebrowzer.QMeasItem = None  # juste pour préconditionner le type et avec acces aux méthodes et attribut de la classe dans le code
        for item in self.qlist.items:
            f = FileReader()
            f.filepath = item.get_filepath()
            trace_nbr = int(item.trace_nbr.itemText(item.trace_nbr.currentIndex()))
            datas = f.get_arrayfromcsv(trace_nbr)
            lim = Meas(pd.DataFrame(datas, columns=['Frequency', 'Field']), item.get_color(), item.get_filename(),
                       item.filepath)
            self.lim_list.append(lim)

    def interp_data(self):
        # Faire une interpolation sur une grille de fréquences commune, même si toutes les meas ont le même ensemble de fréquences
        for i, lim in enumerate(self.lim_list):
            # Créer un nouveau DataFrame avec les valeurs interpolées
            # Utiliser une interpolation logarithmique
            interpolated_values = np.interp(np.log10(self.frequency_axis), np.log10(lim.data['Frequency']),
                                            lim.data['Field'])
            # Remplacer les valeurs en dehors de la plage d'interpolation par une limite tres haute 1000 dBµV/m
            interpolated_values = np.where(
                (self.frequency_axis < lim.data['Frequency'].min()) | (
                        self.frequency_axis > lim.data['Frequency'].max()), 1000,
                interpolated_values)
            # Créer une nouvelle instance de Meas
            self.lim_list[i] = Meas(pd.DataFrame({
                'Frequency': self.frequency_axis,
                'Field': interpolated_values,
            }), lim.color, lim.name, lim.filepath)


class DataProcessor():
    def __init__(self, qamb_list: [filebrowzer.QListPlotItems], qmeas_list: [filebrowzer.QListPlotItems],
                 qlim_list: [filebrowzer.QListPlotItems], root: QWidget):
        self.transducer_factor = 0
        self.root = root
        self.amb_processor = AmbientProcessor(qamb_list, self.root)
        self.meas_processor = MeasProcessor(qmeas_list)
        self.lim_processor = LimitProcessor(qlim_list)
        self.frequency_axis: ndarray = None
        self.align_frequency_datas()
        self.interp_field_value()
        self.amb_processor.merge_ambients()

    def align_frequency_datas(self):
        self.frequency_axis = reduce(np.union1d, (frequencies for frequencies in self.get_allfrequencies()))
        # print(self.frequency_axis)
        self.amb_processor.set_frequency_axis(self.frequency_axis)
        self.meas_processor.set_frequency_axis(self.frequency_axis)
        self.lim_processor.set_frequency_axis(self.frequency_axis)

    def interp_field_value(self):
        self.amb_processor.interp_data()
        self.meas_processor.interp_data()
        self.lim_processor.interp_data()

    def get_allfrequencies(self):
        all_frequencies_list: list(pd.DataFrame) = []
        all_frequencies_list.extend(list(amb.data['Frequency'] for amb in self.amb_processor.ambient_df_list))
        # print(all_frequencies_list[0])
        all_frequencies_list.extend(list(meas.data['Frequency'] for meas in self.meas_processor.meas_list))
        all_frequencies_list.extend(list(lim.data['Frequency'] for lim in self.lim_processor.lim_list))
        return all_frequencies_list

    def get_merged_ambient(self):
        list = []
        if self.amb_processor.merged_ambient is []:
            list = None
        else:
            list = self.amb_processor.merged_ambient
        return list

    def get_meas_list(self):
        list = self.meas_processor.get_meas_list()
        if list is []:
            list = None
        return list

    def get_limit_list(self):
        list = []
        if self.lim_processor.lim_list is []:
            list = None
        else:
            list = self.lim_processor.lim_list
        return list

    def set_transducer_factor(self, transducer_factor):
        """

        :param transducer_factor: in dB
        :return:
        """
        self.transducer_factor = transducer_factor

    def apply_transducer_factor(self):
        """
        Apply the correction factor to the ambient and measurement data
        the correction factor is - transducer factor in dB
        :return:
        """
        for amb in self.amb_processor.ambient_df_list:
            print("debug")
            print(amb.data['Field'][0])
            amb.data['Field'] = amb.data['Field'] - self.transducer_factor
            print(amb.data['Field'][0])
        self.amb_processor.merge_ambients()

        for meas in self.meas_processor.meas_list:
            meas.data['Field'] = meas.data['Field'] - self.transducer_factor



class TableModel(QAbstractTableModel):

    def __init__(self, data: Meas):
        super(TableModel, self).__init__()
        self._data = data

    def data(self, index, role):
        if role == Qt.DisplayRole:
            value = self._data.iloc[index.row(), index.column()]
            return str(value)
            # Gérer la coloration des cellules
        if role == Qt.BackgroundRole:
            ambient_to_limit = self._data.iloc[index.row()]['Ambient to Limit']
            limit_minus_meas = self._data.iloc[index.row()]['Limit'] - self._data.iloc[index.row()]['Measurement']


            if ambient_to_limit < 6:
                return QBrush(QColor(247, 245, 90 ))  # Jaune
            elif limit_minus_meas < 0 and ambient_to_limit > 6:
                return QBrush(QColor(249, 136, 112 ))  # Rouge
        #return QBrush(QColor(255, 255, 255))  # Blanc ou couleur par défaut
        return None

    def rowCount(self, index=QModelIndex()):
        return self._data.shape[0]

    def columnCount(self, index=QModelIndex()):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        # section is the index of the column/row.
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return str(self._data.columns[section])

            if orientation == Qt.Vertical:
                return str(self._data.index[section])


class DataTableProcessor(QWidget):
    def __init__(self, root: QWidget):
        super().__init__()
        self.root = root
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)
        self.analyse_button = QPushButton("show data for current measurement selection")
        self.analyse_button.clicked.connect(lambda: self.setdatas())
        self.root.data_analysis_widget.layout().addWidget(self.analyse_button)
        self.datalist = QTableView()
        self.datamodel = TableModel(Meas(pd.DataFrame(), QColor(0, 0, 0), "init", "none"))
        self.layout.addWidget(self.datalist)
        self.root.data_analysis_widget.layout().addWidget(self.datalist)

    def setdatas(self):
        self.getdatas()

    def getdatas(self):
        merged_ambient_datas: pd.DataFrame = self.root.data_processor.get_merged_ambient().data

        meas_datas_list = self.root.data_processor.get_meas_list()
        selectedindex = self.root.meas_list.get_selected_index()
        #print(selectedindex)
        meas_datas: pd.DataFrame = meas_datas_list[selectedindex].data
        lim_datas: pd.DataFrame = self.root.data_processor.get_limit_list()[0].data
        merged_ambient_datas=merged_ambient_datas.rename(columns={'Field': 'Ambient'})
        meas_datas=meas_datas.rename(columns={'Field': 'Measurement'})
        lim_datas=lim_datas.rename(columns={'Field': 'Limit'})
        merged_datas = merged_ambient_datas.merge(meas_datas, on='Frequency').merge(lim_datas, on='Frequency')
        merged_datas['Ambient to Limit'] = merged_datas['Limit'] - merged_datas['Ambient']
        merged_datas['Meas to Limit'] = merged_datas['Limit'] - merged_datas['Measurement']

        self.setdatalist(merged_datas)
        #print(merged_datas)

    def setdatalist(self, data):
        #print("kesako")
        self.datamodel = TableModel(data)
        self.datalist.setModel(self.datamodel)


class AmbientPostProcessor(QWidget):
    def __init__(self, root: QWidget):
        super().__init__()
        self.peaks_df: pd.DataFrame = None
        self.root = root
        self.pdmodel: pd.PandasModel = None
        # Création du widget principal de l'analyse de l'ambiant
        ambient_analysis_widget = QWidget()

        # Création du layout principal pour le widget d'analyse de l'ambiant
        ambient_analysis_layout = QVBoxLayout(ambient_analysis_widget)

        # Création des widgets pour les paramètres, les résultats et l'analyse post-ambiant
        ambient_analysis_params_widget = QWidget()
        ambient_analysis_results_widget = QWidget()
        ambient_analysis_post_analysis_widget = QWidget()

        # Création des layouts pour les widgets de paramètres, résultats et analyse post-ambiant
        ambient_analysis_params_layout = QHBoxLayout(ambient_analysis_params_widget)
        ambient_analysis_results_layout = QHBoxLayout(ambient_analysis_results_widget)
        ambient_analysis_post_analysis_layout = QHBoxLayout(ambient_analysis_post_analysis_widget)

        # Création et configuration de l'éditeur de ligne pour l'analyse de l'ambiant
        onlyInt = QIntValidator()
        self.ambient_analysis_lineedit = QLineEdit()
        self.ambient_analysis_lineedit.setValidator(onlyInt)
        self.ambient_analysis_lineedit.setText("6")

        # Création du label pour l'analyse de l'ambiant
        ambient_analysis_label = QLabel("discriminator :")

        # Création du bouton pour l'analyse de l'ambiant
        self.ambient_analysis_button = QPushButton("analyse\nde\nl'ambiant")
        self.ambient_analysis_button.clicked.connect(
            lambda: self.root.data_processor.amb_processor.get_merged_ambient_peaks(
                int(round(float(self.ambient_analysis_lineedit.text())))))

        # Ajout du label, de l'éditeur de ligne et du bouton au layout des paramètres
        ambient_analysis_params_layout.addWidget(ambient_analysis_label)
        ambient_analysis_params_layout.addWidget(self.ambient_analysis_lineedit)
        ambient_analysis_params_layout.addWidget(self.ambient_analysis_button)
        # affichage de la table des pics
        self.ambient_analysis_results_tableview = QTableView()
        ambient_analysis_results_layout.addWidget(self.ambient_analysis_results_tableview)

        # Création du bouton pour l'affichage des pics et du bouton pour la couleur dans l'analyse post-ambiant
        self.ambient_post_analysis_plot_button = QPushButton("afficher les peaks")
        self.ambient_post_analysis_plot_button.setEnabled(False)
        self.ambient_post_analysis_plot_button.clicked.connect(self.plot_scatter)
        self.ambient_post_analysis_color = QPushButton(" ")
        self.ambient_post_analysis_color.setStyleSheet("border-radius: 7px;")
        self.button_color: QColor = QColor(92, 85, 33)
        self.ambient_post_analysis_color.clicked.connect(self.open_color_dialog)
        # Ajout des boutons au layout d'analyse post-ambiant
        ambient_analysis_post_analysis_layout.addWidget(self.ambient_post_analysis_plot_button)
        ambient_analysis_post_analysis_layout.addWidget(self.ambient_post_analysis_color)

        # Ajout des widgets de paramètres, de résultats et d'analyse post-ambiant au layout principal
        ambient_analysis_layout.addWidget(ambient_analysis_params_widget)
        ambient_analysis_layout.addWidget(ambient_analysis_results_widget)
        ambient_analysis_layout.addWidget(ambient_analysis_post_analysis_widget)

        # Ajout du widget d'analyse de l'ambiant au layout principal de l'application
        self.root.ambiant_analysis_widget.layout().addWidget(ambient_analysis_widget)

    def set_ambient_tableview_model(self):
        print("ok")
        self.pdmodel = PandasModel(self.peaks_df)
        self.ambient_analysis_results_tableview.setModel(self.pdmodel)
        self.ambient_post_analysis_plot_button.setEnabled(True)

    def set_peaks_df(self, df: pd.DataFrame):
        self.peaks_df = df.drop(['Lower Base Level', 'Upper Base Level'], axis=1)
        self.set_ambient_tableview_model()

    def open_color_dialog(self):
        self.color_dialog = QColorDialog(self)
        self.color_dialog.colorSelected.connect(self.change_button_color)
        self.color_dialog.show()

    def change_button_color(self, color: QColor):
        self.ambient_post_analysis_color.setStyleSheet(f"border-radius: 7px;background-color: {color.name()};")
        self.button_color = color

    def plot_scatter(self):
        # Récupération du DataFrame des pics
        peaks_df = self.peaks_df
        # Récupération du DataFrame d'origine
        original_df = self.root.data_processor.get_merged_ambient().data  # Remplacez par votre fonction appropriée

        if peaks_df is not None and original_df is not None:
            peaksubsets = []  # Liste pour stocker tous les sous-ensembles
            lowsubsets = []
            upsubsets = []
            # Pour chaque pic, récupérer la partie correspondante du DataFrame original
            for _, row in peaks_df.iterrows():
                lower_index = int(row['Lower Base Index'])
                peak_index = int(row['Peak Index'])
                upper_index = int(row['Upper Base Index'])
                # print(lower_index)
                # upper_index = int(row['Upper Base Index'])
                # print(upper_index)
                low_subset_df = original_df.iloc[lower_index]
                peak_subset_df = original_df.iloc[peak_index]
                up_subset_df = original_df.iloc[upper_index]
                # subset_df = original_df.iloc[lower_index:upper_index + 1]  # +1 pour inclure la borne supérieure
                lowsubsets.append(low_subset_df)
                peaksubsets.append(peak_subset_df)
                upsubsets.append(up_subset_df)

            # Combiner tous les sous-ensembles en un seul DataFrame
            combined_low_df = pd.concat(lowsubsets)
            combined_peak_df = pd.concat(peaksubsets)
            combined_up_df = pd.concat(upsubsets)
            print(self.button_color.getRgb())
            x_low = combined_low_df['Frequency'].values
            y_low = combined_low_df['Field'].values
            self.root.graphWidget.plot(x_low, y_low, pen=None, symbol='t', symbolBrush=self.button_color.getRgb())
            x_peak = combined_peak_df['Frequency'].values
            y_peak = combined_peak_df['Field'].values
            self.root.graphWidget.plot(x_peak, y_peak, pen=None, symbol='o', symbolBrush=self.button_color.getRgb())
            x_up = combined_up_df['Frequency'].values
            y_up = combined_up_df['Field'].values
            self.root.graphWidget.plot(x_up, y_up, pen=None, symbol='+', symbolBrush=self.button_color.getRgb())


class PandasModel(QAbstractTableModel):
    def __init__(self, df=pd.DataFrame(), parent=None):
        QAbstractTableModel.__init__(self, parent)
        self._df = df

    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role != Qt.DisplayRole:
            return None
        if orientation == Qt.Horizontal:
            try:
                return self._df.columns.tolist()[section]
            except IndexError:
                return None
        elif orientation == Qt.Vertical:
            try:
                return self._df.index.tolist()[section]
            except IndexError:
                return None

    def data(self, index, role=Qt.DisplayRole):
        if role != Qt.DisplayRole:
            return None
        if not index.isValid():
            return None
        return str(self._df.iloc[index.row(), index.column()])

    def rowCount(self, parent=None):
        return len(self._df.index)

    def columnCount(self, parent=None):
        return len(self._df.columns)


class MeasPostProcessor(QWidget):
    def __init__(self, root: QWidget):
        super().__init__()
        self.model = QStandardItemModel()
        self.root = root
        self.ambient_merged = None
        self.exceedings_list: [pd.DataFrame] = None
        self.meas_list = None
        self.limits_list = None
        self.pdmodel: pd.PandasModel = None
        # Création du widget principal de l'analyse des mesures
        meas_analysis_widget = QWidget()
        # Création du layout principal pour le widget d'analyse de la mesure
        meas_analysis_layout = QVBoxLayout(meas_analysis_widget)

        # Création des widgets pour les paramètres, les résultats et l'analyse post-mesure
        meas_analysis_params_widget = QWidget()
        meas_analysis_results_widget = QWidget()
        meas_analysis_post_analysis_widget = QWidget()

        # Création des layouts pour les widgets de paramètres, résultats et analyse post-mesure
        meas_analysis_params_layout = QHBoxLayout(meas_analysis_params_widget)
        meas_analysis_results_layout = QHBoxLayout(meas_analysis_results_widget)
        meas_analysis_post_analysis_layout = QHBoxLayout(meas_analysis_post_analysis_widget)

        # Création et configuration de l'éditeur de ligne pour l'analyse des mesures en fonction de la limite d'attention
        onlyInt = QIntValidator()
        self.meas_analysis_lineedit = QLineEdit()
        self.meas_analysis_lineedit.setValidator(onlyInt)
        self.meas_analysis_lineedit.setText("6")

        # Création du label pour l'analyse de l'ambiant
        meas_analysis_label = QLabel("attention factor (dB) :")

        # Création du bouton pour l'analyse de l'ambiant
        # self.meas_analysis_button = QPushButton("Measurements\nanalysis")
        self.meas_post_analysis_plot_button = QPushButton("Find exceeding values")
        self.meas_post_analysis_plot_button.setEnabled(False)
        self.ambient_post_analysis_color = QPushButton(" ")
        self.ambient_post_analysis_color.setStyleSheet("border-radius: 7px;")
        self.button_color: QColor = QColor(92, 85, 33)
        self.ambient_post_analysis_color.clicked.connect(self.open_color_dialog)
        self.meas_post_analysis_plot_button.clicked.connect(self.process_exceedings)

        #        self.meas_analysis_button.clicked.connect(
        #            lambda: self.root.data_processor.amb_processor.get_merged_ambient_peaks(
        #                int(round(float(self.meas_analysis_lineedit.text())))))

        # Ajout du label, de l'éditeur de ligne et du bouton au layout des paramètres
        meas_analysis_post_analysis_layout.addWidget(meas_analysis_label)
        meas_analysis_post_analysis_layout.addWidget(self.meas_analysis_lineedit)
        # meas_analysis_params_layout.addWidget(self.meas_post_analysis_plot_button)
        # meas_analysis_params_layout.addWidget(self.meas_post_analysis_plot_button)
        # affichage de la table des pics
        self.meas_analysis_results_treeview = QTreeView()
        meas_analysis_results_layout.addWidget(self.meas_analysis_results_treeview)

        # Création du bouton pour l'affichage des pics et du bouton pour la couleur dans l'analyse post-ambiant

        # Ajout des boutons au layout d'analyse post-ambiant
        meas_analysis_post_analysis_layout.addWidget(self.meas_post_analysis_plot_button)
        meas_analysis_post_analysis_layout.addWidget(self.ambient_post_analysis_color)

        # Ajout des widgets de paramètres, de résultats et d'analyse post-ambiant au layout principal
        meas_analysis_layout.addWidget(meas_analysis_post_analysis_widget)
        meas_analysis_layout.addWidget(meas_analysis_results_widget)

        meas_analysis_report_widget = QWidget()
        meas_analysis_report_widget.setLayout(QHBoxLayout())
        meas_analysis_report_stat_label = QLabel("Status")
        self.status_combobox = QComboBox()
        self.status_combobox.addItems([status.value for status in Status])
        self.status_combobox.setStyleSheet("QComboBox { color: black; }")

        self.argument_combobox = QComboBox()
        self.argument_combobox.addItems([argument.value for argument in Argument])
        self.argument_combobox.setStyleSheet("QComboBox { color: black; }")

        self.ok_button = QPushButton("Valider l'argument")
        self.ok_button.clicked.connect(self.update_selected_row)

        # Ajoutez ces widgets à votre layout

        meas_analysis_report_widget.layout().addWidget(self.status_combobox)
        meas_analysis_report_widget.layout().addWidget(self.argument_combobox)
        meas_analysis_report_widget.layout().addWidget(self.ok_button)
        meas_analysis_layout.addWidget(meas_analysis_report_widget)
        # element pour génération de rapport

        self.generate_report_button = QPushButton("Generate Report")
        self.generate_report_button.clicked.connect(self.generate_report)
        meas_analysis_layout.addWidget(self.generate_report_button)

        # Ajout du widget d'analyse de l'ambiant au layout principal de l'application
        self.root.analysis_widget.layout().addWidget(meas_analysis_widget)

    def generate_report(self):
        document = Document()

        # Ajout du titre
        document.add_heading('Résultats', level=2)

        # Ajout du tableau
        document.add_heading(f"Table des dépassements", level=3)
        table = document.add_table(rows=1, cols=7)
        table.autofit = True
        table.style = 'Table Grid'
        headers = ['Mesure', 'Dépassement\nN°', 'fmin\n(MHz)', 'fmax\n(MHz)', 'Dépassement\nMax\n(dB)', 'Arguments',
                   'Status']
        hdr_cells = table.rows[0].cells
        for hdr_cell, header_text in zip(hdr_cells, headers):
            paragraph = hdr_cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = header_text
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Récupération des données du TreeView
        for row in range(self.model.rowCount()):
            parent = self.model.item(row)
            for child_row in range(parent.rowCount()):
                name = parent.text()
                group = parent.child(child_row, 1).text()
                bookmark_name = f"detail_{name}_{group}"
                status_item = parent.child(child_row, 4).text()
                freqmin = str(round(float(parent.child(child_row, 2).text()) / 1E6, 3))
                freqmax = str(round(float(parent.child(child_row, 3).text()) / 1E6, 3))
                maxexceeding = str(round(float(parent.child(child_row, 6).text()), 1))
                argument = parent.child(child_row, 5).text()

                row_cells = table.add_row().cells

                p = row_cells[1].paragraphs[0]

                row_cells[0].text = name
                # row_cells[1].text = group
                bookmark_name = f"detail_{name}_{group}"
                self.add_hyperlink(p, bookmark_name)  # Add hyperlink to the group cell
                row_cells[2].text = freqmin
                row_cells[3].text = freqmax
                row_cells[4].text = maxexceeding
                row_cells[5].text = argument
                row_cells[6].text = status_item

                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrer le contenu de chaque cellule

        # Récupération des noms des courbes 'overview' à inclure
        meas_names_to_include = set()
        for row in range(self.model.rowCount()):
            parent = self.model.item(row)
            meas_name = parent.text()
            document.add_heading(f"Mesure {meas_name}", level=3)
            # Masquer les courbes non désirées
            plot_widget = self.root.graphWidget

            plot_item = plot_widget.getPlotItem()
            plot_widget.setLogMode(x=True, y=False)
            plot_item.getViewBox().autoRange()

            legend = plot_item.legend
            data_items = plot_item.listDataItems()

            # Un dictionnaire pour conserver la visibilité originale des courbes
            original_visibility = {}

            for item in data_items:
                curve_name = item.opts['name']
                print(f"plot curve name :{curve_name}")
                print(f"meas_names_to_include :{meas_name}")

                if meas_name not in curve_name and "lim" not in curve_name and "Max Ambient" not in curve_name and "50121" not in curve_name:
                    original_visibility[item] = item.isVisible()
                    item.setVisible(False)

                    # Masquer l'élément correspondant dans la légende
                    for sample, label in legend.items:
                        if label.text == curve_name:
                            sample.hide()
                            label.hide()

            # Exporter le graphique
            image_path = 'plot.png'
            exporter = ImageExporter(plot_item)
            exporter.export(image_path)

            # Ajout de l'image au document Word
            document.add_picture(image_path)

            # Restaurer la visibilité originale des courbes et des éléments de la légende
            for item, was_visible in original_visibility.items():
                item.setVisible(was_visible)

                # Restaurer la visibilité dans la légende
                for sample, label in legend.items:
                    if label.text == item.opts['name']:
                        sample.show()
                        label.show()

        # Crée un tableau pour organiser les bookmarks et les images
        document.add_heading(f"Détails des dépassements", level=3)
        table = document.add_table(rows=0, cols=2)

        for row in range(self.model.rowCount()):
            parent = self.model.item(row)
            meas_name = parent.text()
            document.add_heading(f"Mesure {meas_name}", level=2)
            for child_row in range(0, parent.rowCount(), 2):
                #         group = parent.child(child_row, 1).text()
                #         freqmin = float(parent.child(child_row, 2).text())
                #         freqmax = float(parent.child(child_row, 3).text())
                #         zoom_shift = 1E6
                #         # Masquer les courbes non désirées
                #         plot_widget: pyqtgraph.PlotWidget = self.root.graphWidget
                #
                #         plot_item = plot_widget.getPlotItem()
                #         plot_widget.setLogMode(x=False, y=False)
                #         plot_widget.setXRange(freqmin - zoom_shift, freqmax + zoom_shift)
                # Ajoute une nouvelle ligne au tableau
                row_cells = table.add_row().cells

                for col in range(2):
                    if child_row + col < parent.rowCount():
                        group = parent.child(child_row + col, 1).text()
                        freqmin = float(parent.child(child_row + col, 2).text())
                        freqmax = float(parent.child(child_row + col, 3).text())
                        zoom_shift = 1E6
                        bookmark_name = f"detail_{meas_name}_{group}"
                        legend = plot_item.legend
                        data_items = plot_item.listDataItems()

                        # Un dictionnaire pour conserver la visibilité originale des courbes
                        original_visibility = {}
                        #
                        for item in data_items:
                            curve_name = item.opts['name']
                            #             print(f"plot curve name :{curve_name}")
                            #             print(f"meas_names_to_include :{meas_name}")
                            #
                            if meas_name not in curve_name and "lim" not in curve_name and "Max Ambient" not in curve_name and "50121" not in curve_name:
                                original_visibility[item] = item.isVisible()
                                item.setVisible(False)

                                # Masquer l'élément correspondant dans la légende
                                for sample, label in legend.items:
                                    if label.text == curve_name:
                                        sample.hide()
                                        label.hide()

                        # Ajoute le bookmark à la cellule
                        p = row_cells[col].paragraphs[0]
                        self.add_bookmark(p, bookmark_name, group, bookmark_name)
                        plot_widget: pyqtgraph.PlotWidget = self.root.graphWidget
                        plot_widget.setLogMode(x=False, y=False)
                        plot_widget.setXRange(freqmin - zoom_shift, freqmax + zoom_shift)

                        # Exporte et ajoute l'image à la cellule
                        image_path = 'plot.png'
                        inches_width = 7 / 2.54
                        exporter = ImageExporter(plot_item)
                        exporter.export(image_path)

                        row_cells[col].add_paragraph().add_run().add_picture(image_path, width=Inches(inches_width))

                        # Restaurer la visibilité originale des courbes et des éléments de la légende
                        for item, was_visible in original_visibility.items():
                            item.setVisible(was_visible)

                            # Restaurer la visibilité dans la légende
                            for sample, label in legend.items:
                                if label.text == item.opts['name']:
                                    sample.show()
                                    label.show()
        plot_widget = self.root.graphWidget
        plot_widget.setLogMode(x=True, y=False)
        plot_item.getViewBox().autoRange()
        document.save('rapport.docx')
        QMessageBox.information(self, "Rapport généré", "Le rapport a été généré avec succès.")

    def add_hyperlink(self, paragraph, bookmark_name):
        run = paragraph.add_run()
        fldCharBegin = OxmlElement('w:fldChar')
        fldCharBegin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f"REF {bookmark_name} \\h"

        fldCharSeparate = OxmlElement('w:fldChar')
        fldCharSeparate.set(qn('w:fldCharType'), 'separate')

        hyperlink_text = OxmlElement('w:r')
        hyperlink_t = OxmlElement('w:t')
        hyperlink_t.text = bookmark_name  # This will be the displayed text for the hyperlink
        hyperlink_text.append(hyperlink_t)

        fldCharEnd = OxmlElement('w:fldChar')
        fldCharEnd.set(qn('w:fldCharType'), 'end')

        run._r.append(fldCharBegin)
        run._r.append(instrText)
        run._r.append(fldCharSeparate)
        run._r.append(hyperlink_text)
        run._r.append(fldCharEnd)

        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color

        return run

    def add_bookmark(self, paragraph, bookmark_name, group_id, visible_text):
        run = paragraph.add_run()

        bookmarkStart = OxmlElement('w:bookmarkStart')
        bookmarkStart.set(qn('w:id'), str(group_id))
        bookmarkStart.set(qn('w:name'), bookmark_name)

        run._r.append(bookmarkStart)

        run = paragraph.add_run(visible_text)  # Ajout du texte visible
        run.bold = True  # Rendre le texte en gras
        bookmarkEnd = OxmlElement('w:bookmarkEnd')
        bookmarkEnd.set(qn('w:id'), str(group_id))

        run._r.append(bookmarkEnd)

        return run

    def set_meas_tableview_model(self):
        # print("ok")
        self.pdmodel = PandasModel(self.exeedings_df)
        self.meas_analysis_results_tableview.setModel(self.pdmodel)

    def set_exceedings_df(self, df: pd.DataFrame):
        self.exeedings_df = df.drop(['Lower Base Level', 'Upper Base Level'], axis=1)
        self.set_meas_tableview_model()

    def open_color_dialog(self):
        self.color_dialog = QColorDialog(self)
        self.color_dialog.colorSelected.connect(self.change_button_color)
        self.color_dialog.show()

    def change_button_color(self, color: QColor):
        self.meas_post_analysis_color.setStyleSheet(f"border-radius: 7px;background-color: {color.name()};")
        self.button_color = color

    def process_exceedings(self):
        self.ambient_merged = self.root.data_processor.get_merged_ambient().data
        self.meas_list = self.root.data_processor.get_meas_list()
        self.limits_list = self.root.data_processor.get_limit_list()
        # Récupération du DataFrame des pics
        exeedings_df: pd.DataFrame = None
        self.exceedings_list: [pd.DataFrame] = []
        ambient_merged_df: pd.DataFrame = self.ambient_merged
        filtered_measures_df: pd.DataFrame = None
        attention_limit = int(round(float(self.meas_analysis_lineedit.text())))
        self.exceedings_list.clear()
        alldataframes: pd.DataFrame = None
        for limit in self.limits_list:
            limit_df = limit.data.copy()
            for meas in self.meas_list:
                meas_df = meas.data.copy()
                alldataframes = self.ambient_merged.copy()
                alldataframes.rename(columns={'Field': 'Max Ambient'}, inplace=True)
                alldataframes = alldataframes.merge(meas_df, on='Frequency')
                alldataframes.rename(columns={'Field': 'Measure'}, inplace=True)
                alldataframes = alldataframes.merge(limit_df, on='Frequency')
                alldataframes.rename(columns={'Field': 'Limit'}, inplace=True)

                # Filtering the measurements where the ambient value is less than the limit value - the attention limit
                filtered_measures_df = alldataframes.loc[
                    alldataframes['Max Ambient'] < alldataframes['Limit'] - attention_limit]

                # Identifying exceedings
                exceedings_df: pd.DataFrame = filtered_measures_df.loc[
                    filtered_measures_df['Measure'] > filtered_measures_df['Limit']]

                # pd.set_option('display.max_rows', None)
                # print(exceedings_df)
                grouped_exceedings_df: pd.DataFrame = exceedings_df.copy()
                grouped_exceedings_df['Group'] = (exceedings_df.copy().index.to_series().diff() != 1).cumsum()
                # print(f"max : {max_exceeding}")
                # Appending the exceedings to the list
                self.exceedings_list.append((meas.name, grouped_exceedings_df))

        self.plot_scatter()
        self.fill_tree_view()

    def plot_scatter(self):

        # Récupération du DataFrame d'origine
        for name, exceedings_df in self.exceedings_list:

            if exceedings_df is not None:
                # print(self.button_color.getRgb())
                x_low = exceedings_df['Frequency'].values
                y_low = exceedings_df['Measure'].values
                self.root.graphWidget.plot(x_low, y_low, pen=None, symbol='x', symbolBrush=self.button_color.getRgb(),
                                           name=f"{name} exceeding values")

    def fill_tree_view(self):
        # Création du modèle
        self.model = QStandardItemModel()
        self.model.setHorizontalHeaderLabels(
            ['Name', 'Group', 'Min Frequency', 'Max Frequency', 'Status', 'Arguments', 'Max Exceeding'])

        # Ajout des données
        for elem_name, grouped_exceedings_df in self.exceedings_list:
            # Création de l'élément parent
            parent = QStandardItem(elem_name)
            self.model.appendRow(parent)

            # Regroupement par 'Group' et calcul des fréquences minimale et maximale
            group_df = grouped_exceedings_df.groupby('Group')['Frequency'].agg(['min', 'max'])
            grouped_exceedings_df['Exceeding'] = (
                    grouped_exceedings_df['Measure'] - grouped_exceedings_df['Limit']).round(3)
            max_exceeding_per_group = grouped_exceedings_df.groupby('Group')['Exceeding'].max().reset_index()
            group_df = group_df.merge(max_exceeding_per_group, on='Group')
            # Ajout des groupes comme enfants de l'élément parent
            for group, row in group_df.iterrows():
                # print(group)

                meas_name = QStandardItem(str(elem_name))
                group_item = QStandardItem(str(group))
                min_freq_item = QStandardItem(str(row['min']))
                max_freq_item = QStandardItem(str(row['max']))
                status_item = QStandardItem(Status.NOK.value)
                argument_item = QStandardItem(Argument.NOK.value)
                max_exceeding_item = QStandardItem(str(row['Exceeding']))
                parent.appendRow([meas_name, group_item, min_freq_item, max_freq_item, status_item, argument_item,
                                  max_exceeding_item])

        # Mettre à jour le modèle dans la vue
        self.meas_analysis_results_treeview.setModel(self.model)
        self.meas_analysis_results_treeview.setExpandsOnDoubleClick(True)
        for column in range(self.meas_analysis_results_treeview.model().columnCount()):
            self.meas_analysis_results_treeview.resizeColumnToContents(column)
        # spinbox = QSpinBox()
        # spinbox.setMinimum(0)
        # spinbox.setMaximum(100)
        # Replace 0, 0 with the actual row and column of the cell you want to add the spinbox to
        # self.meas_analysis_results_treeview.setIndexWidget(model.index(0, 0), spinbox)
        self.meas_analysis_results_treeview.clicked.connect(self.on_item_clicked)
        self.update_colors()

    def on_item_clicked(self, index):
        # Obtenez le modèle de l'index
        model = index.model()

        # Obtenez les informations requises
        name = model.itemFromIndex(index.sibling(index.row(), 0)).text()
        min_freq = float(model.itemFromIndex(index.sibling(index.row(), 2)).text())
        max_freq = float(model.itemFromIndex(index.sibling(index.row(), 3)).text())

        self.root.graphWidget.setLogMode(x=False, y=False)
        self.root.graphWidget.setRange(xRange=[(min_freq - 1E6), (max_freq + 1E6)])

        for item in self.root.meas_list.items:
            if item.get_filename() == name:
                self.root.meas_list.select_item(item)
                break
        self.root.data_table_processor.setdatas()
        row_to_scroll = None
        for row in range(self.root.data_table_processor.datalist.model().rowCount()):
            freq_value = self.root.data_table_processor.datalist.model().data(
                self.root.data_table_processor.datalist.model().index(row, 0), Qt.DisplayRole)

            #freq_value = self.root.data_table_processor.datalist.model().data(self.root.data_table_processor.datalist.model().index(row,0))  # 0 est l'index de la colonne où se trouve la fréquence
            if float(freq_value) >= min_freq:
                row_to_scroll = row
                break

        if row_to_scroll is not None:
            target_index = self.root.data_table_processor.datalist.model().index(row_to_scroll,
                                                       0)  # L'index de la ligne à faire défiler, première colonne
            self.root.data_table_processor.datalist.scrollTo(target_index,
                                   QAbstractItemView.PositionAtTop)  # Fait défiler la vue jusqu'à l'index trouvé

    def update_selected_row(self):
        selected_indexes = self.meas_analysis_results_treeview.selectedIndexes()
        if selected_indexes:
            index = selected_indexes[0]
            model = index.model()
            status_item = model.itemFromIndex(index.sibling(index.row(), 4))
            argument_item = model.itemFromIndex(index.sibling(index.row(), 5))
            status_item.setText(self.status_combobox.currentText())
            argument_item.setText(self.argument_combobox.currentText())
            self.update_colors()

    def update_colors(self):
        red_color = QBrush(QColor(255, 0, 0))
        green_color = QBrush(QColor(0, 255, 0))

        for row in range(self.model.rowCount()):
            parent = self.model.item(row)
            has_nok_child = False

            # Traverse all children of the current parent
            for child_row in range(parent.rowCount()):
                status_item = parent.child(child_row, 4)  # Access the "status" item in column 4 of the child

                # Use the "status" item
                if status_item:
                    status_text = status_item.text()
                    color = red_color if status_text == Status.NOK.value else green_color
                    for col in range(7):  # Assuming you have 6 columns
                        self.model.setData(self.model.index(child_row, col, parent.index()), color, Qt.ForegroundRole)
                    if status_text == Status.NOK.value:
                        has_nok_child = True

            if has_nok_child:
                for col in range(7):  # Assuming you have 6 columns
                    self.model.setData(self.model.index(row, col), red_color, Qt.ForegroundRole)
            if not has_nok_child:
                for col in range(7):  # Assuming you have 6 columns
                    self.model.setData(self.model.index(row, col), green_color, Qt.ForegroundRole)


class Status(Enum):
    OK = "OK"
    NOK = "NOK"


class Argument(Enum):
    AmbientPresence = "Present in Ambient"
    DataChanelSideEffect = "Channel Side effect"
    Transient = "Caused by ambient transient"
    QpeakCompliant = "Measured in Qpeak and ok"
    HumanActivity = "Due to Human activity"
    NOK = "exceedinvalue to be investigated"
